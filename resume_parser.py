import pdfplumber
from docx import Document
from io import BytesIO

def extract_text_from_pdf(file_obj):
    """Extract text from an uploaded PDF file while preserving formatting."""
    with pdfplumber.open(BytesIO(file_obj.read())) as pdf:
        extracted_text = []
        for page in pdf.pages:
            # Get text with layout to preserve spacing
            text = page.extract_text(layout=True)
            if text:
                extracted_text.append(text)
        return "\n".join(extracted_text)

def extract_text_from_docx(file_obj):
    """Extract formatted text from a DOCX file while preserving structure."""
    doc = Document(BytesIO(file_obj.read()))
    extracted_text = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Preserve bullet points
            if para.style.name.startswith("List") or para.text.startswith("-"):
                text = f"• {text}"
            
            # Check for bold text and highlight it
            if any(run.bold for run in para.runs):
                text = f"**{text}**"  # Markdown-style bold
            
            extracted_text.append(text)

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                extracted_text.append(row_text)
    
    return "\n".join(extracted_text)

def extract_text(file_obj):
    """Detect file type and extract text while preserving formatting."""
    if file_obj.name.endswith(".pdf"):
        return extract_text_from_pdf(file_obj)
    elif file_obj.name.endswith(".docx"):
        return extract_text_from_docx(file_obj)
    else:
        raise ValueError("Unsupported file format. Please upload a PDF or DOCX.")

